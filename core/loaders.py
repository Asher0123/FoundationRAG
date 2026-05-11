from pypdf import PdfReader
from docx import Document as DocxDocument
from typing import List,Any,Dict
from abc import ABC, abstractmethod
from dataclasses import dataclass,field
from exceptions import LoaderError
import os
from pathlib import Path

@dataclass
class Document:
    """A dataclass which contains content and metadata."""
    content: str
    metadata: Dict[str, Any]=field(default_factory=dict)

class BaseLoader(ABC):

    @abstractmethod
    def load_document(self, path: str) -> List[Document]:
        pass

class PDFLoader(BaseLoader):
    """
    Loader for PDF documents.

    Extracts text page-by-page from a PDF file
    and returns a list of Document objects with metadata.

    Uses pypdf's PdfReader.
    """
    def load_document(self, path: str) -> List[Document]:
        """
        Load and extract text from a PDF document.

        Args:
            path: Path to the PDF document.

        Returns:
            List of Document objects containing:
            - extracted page text as content
            - metadata such as page number, source path,
              and file type.
        """
        try:
            docs=[]
            if not path.lower().endswith('.pdf'):
                raise LoaderError(f"Use PDF file.")
            if not os.path.exists(path):
                raise LoaderError(f"Provided path is incorrect: {path}")
            if path.lower().endswith('.pdf'):
                document=PdfReader(path)
                for page_num,page in enumerate(document.pages):
                    doc=Document(content=page.extract_text() or "", metadata={'page_no':page_num,'source':Path(path).name,'file_type':'PDF'})
                    docs.append(doc)
                return docs
            else:
                raise LoaderError(f"Provide a pdf file as input.")                   
        except Exception as e:
            raise LoaderError(
                f"Failed to load PDF: {path} due to {e}"
            ) from e
        
    

class DOCXLoader(BaseLoader):
    """
    Loader for DOCX documents.

    Extracts paragraph and table content from a DOCX file
    and returns a list of Document objects with metadata.

    Uses python-docx's Document parser.
    """
    def load_document(self, path: str) -> List[Document]:
        """
        Load and extract text from a DOCX document.

        Args:
            path: Path to the DOCX document.

        Returns:
            List of Document objects containing:
            - extracted document text as content
            - metadata such as source path and file type.
        """
        try:
            if not path.lower().endswith('.docx'):
                raise LoaderError(f"Use DOCX file.")
            if not os.path.exists(path):
                raise LoaderError(f"Provided path is incorrect: {path}")
            if path.lower().endswith('.docx'):
                document=DocxDocument(path)
                text=""
                paragraphs = [para.text.strip() for para in document.paragraphs if para.text.strip()]
                text = "\n".join(paragraphs)

                table_texts = []
                for table in document.tables:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        if len(cells) >= 2:
                            kv_pairs = [f"{cells[i]}: {cells[i + 1]}" for i in range(0, len(cells) - 1, 2)]
                            table_texts.append("; ".join(kv_pairs))
                        else:
                            table_texts.append(" | ".join(cells)) 

                full_text=text+"\n\n"+"\n".join(table_texts)
                return [Document(full_text,{'source':Path(path).name,'file_type':'DOCX'})]
            else:
                raise LoaderError(f"Provide a docx file as input.")                          
        except Exception as e:
            raise LoaderError(
                f"Failed to load DOCX: {path}"
            ) from e        